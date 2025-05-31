from typing import Dict, Any, List, Optional
import os
from pathlib import Path

class DocumentExtractor:
    """
    Document content extractor for various file formats.
    """
    
    name = "document_extractor"
    description = "Extracts content from various document formats"
    
    def __init__(self, upload_dir: Optional[Path] = None):
        """
        Initialize the document extractor.
        
        Args:
            upload_dir: Directory where uploaded files are stored
        """
        self.upload_dir = upload_dir
    
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text content from a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary with extracted content and metadata
        """
        if self.upload_dir and not os.path.isabs(file_path):
            file_path = self.upload_dir / file_path
        
        if not os.path.exists(file_path):
            return {'error': f'File not found: {file_path}', 'success': False}
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf(file_path)
            elif file_ext in ['.txt', '.md']:
                return self._extract_text_file(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_word_document(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_excel_document(file_path)
            else:
                return {
                    'error': f'Unsupported file type: {file_ext}',
                    'success': False
                }
        except Exception as e:
            return {
                'error': f'Failed to extract content: {str(e)}',
                'success': False
            }
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from PDF files.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted content
        """
        try:
            import pdfplumber
            
            text_content = []
            metadata = {
                'file_path': file_path,
                'file_size': os.path.getsize(file_path),
                'num_pages': 0
            }
            
            with pdfplumber.open(file_path) as pdf:
                metadata['num_pages'] = len(pdf.pages)
                
                # Extract metadata if available
                if pdf.metadata:
                    metadata.update({
                        'title': pdf.metadata.get('Title'),
                        'author': pdf.metadata.get('Author'),
                        'creation_date': pdf.metadata.get('CreationDate'),
                        'modification_date': pdf.metadata.get('ModDate')
                    })
                
                # Extract text from each page
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num} ---\n{page_text}")
            
            full_content = "\n\n".join(text_content)
            
            return {
                'success': True,
                'content': full_content,
                'file_type': 'pdf',
                'num_pages': metadata['num_pages'],
                'metadata': metadata
            }
            
        except ImportError:
            return {
                'error': 'pdfplumber library not available. Please install it with: pip install pdfplumber',
                'success': False
            }
        except Exception as e:
            return {
                'error': f'Failed to extract PDF content: {str(e)}',
                'success': False
            }
    
    def _extract_text_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from text files.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extracted content
        """
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            return {
                'success': True,
                'content': content,
                'file_type': 'text',
                'content_length': len(content),
                'metadata': {
                    'file_path': file_path,
                    'file_size': os.path.getsize(file_path),
                    'encoding': 'utf-8'
                }
            }
        except Exception as e:
            return {
                'error': f'Failed to read text file: {str(e)}',
                'success': False
            }
    
    def _extract_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from Word documents.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Extracted content
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(" | ".join(row_data))
                if table_data:
                    tables_text.append("\n".join(table_data))
            
            # Combine content
            content_parts = []
            if paragraphs:
                content_parts.append("\n".join(paragraphs))
            if tables_text:
                content_parts.append("\n--- Tables ---\n" + "\n\n".join(tables_text))
            
            full_content = "\n\n".join(content_parts)
            
            return {
                'success': True,
                'content': full_content,
                'file_type': 'word',
                'num_paragraphs': len(paragraphs),
                'num_tables': len(doc.tables),
                'metadata': {
                    'file_path': file_path,
                    'file_size': os.path.getsize(file_path)
                }
            }
            
        except ImportError:
            return {
                'error': 'python-docx library not available. Please install it with: pip install python-docx',
                'success': False
            }
        except Exception as e:
            return {
                'error': f'Failed to extract Word document content: {str(e)}',
                'success': False
            }
    
    def _extract_excel_document(self, file_path: str) -> Dict[str, Any]:
        """
        Extract data from Excel documents.
        
        Args:
            file_path: Path to the Excel document
            
        Returns:
            Extracted content
        """
        try:
            import pandas as pd
            
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            content_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert to text representation
                sheet_text = f"--- Sheet: {sheet_name} ---\n"
                sheet_text += df.to_string(index=False)
                content_parts.append(sheet_text)
                
                sheets_data[sheet_name] = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist()
                }
            
            full_content = "\n\n".join(content_parts)
            
            return {
                'success': True,
                'content': full_content,
                'file_type': 'excel',
                'num_sheets': len(excel_file.sheet_names),
                'sheets_info': sheets_data,
                'metadata': {
                    'file_path': file_path,
                    'file_size': os.path.getsize(file_path),
                    'sheet_names': excel_file.sheet_names
                }
            }
            
        except ImportError:
            return {
                'error': 'pandas or openpyxl library not available. Please install with: pip install pandas openpyxl',
                'success': False
            }
        except Exception as e:
            return {
                'error': f'Failed to extract Excel content: {str(e)}',
                'success': False
            }
